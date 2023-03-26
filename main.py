import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from rake_nltk import Rake
import openai
import os

openai.api_key = os.getenv("OPENAI_API_KEY")

def extract_keywords(text):
    rake = Rake()
    rake.extract_keywords_from_text(text)
    return rake.get_ranked_phrases()

def update_paragraphs(paragraphs, new_content):
    new_content_lines = new_content.split("\n")
    for i, para in enumerate(paragraphs):
        if i < len(new_content_lines):
            para.text = new_content_lines[i]
            para.style = para.style
        else:
            para.text = ""

def generate_resume(base_resume, job_description, keywords):
    base_resume_text = '\n'.join([para.text for para in base_resume.paragraphs])
    messages = [
        {"role": "system", "content": "You are an AI that helps users optimize their resumes for Applicant Tracking Systems (ATS) by incorporating relevant keywords and tailoring the content to specific job descriptions. Your goal is to rewrite the user's resume to increase their chances of getting noticed by the ATS and the hiring manager."},
        {"role": "user", "content": f"Here is the job description:\n\n{job_description}. Keywords:\n\n{', '.join(keywords)}.\n\nPlease optimize my resume for this job."},
        {"role": "user", "content": f"Here is my current resume:\n\n{base_resume_text}.\n\nPlease provide only the optimized resume for this job."}
    ]
    completion = openai.ChatCompletion.create(model="gpt-4", messages=messages)
    return completion.choices[0].message["content"]

def generate_cover_letter(job_description, rewritten_resume, keywords):
    rewritten_resume_text = '\n'.join([para.text for para in rewritten_resume.paragraphs])
    messages = [
        {"role": "system", "content": "You are an AI that helps users create tailored cover letters for specific job descriptions, incorporating relevant keywords and highlighting their skills and experiences. Your goal is to write a compelling cover letter that demonstrates the user's interest in the position and showcases their qualifications."},
        {"role": "user", "content": f"Here is the job description: {job_description}. Keywords: {', '.join(keywords)}. Please write a cover letter for this job based on my optimized resume."},
        {"role": "user", "content": f"Here is my optimized resume:\n\n{rewritten_resume_text}"}
    ]
    completion = openai.ChatCompletion.create(model="gpt-4", messages=messages)
    return completion.choices[0].message["content"]

def upload_base_resume():
    file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if file_path:
        base_resume_path.set(file_path)

def generate_output():
    base_resume = Document(base_resume_path.get())
    job_desc = job_description_text.get("1.0", tk.END).strip()

    if not base_resume or not job_desc:
        messagebox.showerror("Error", "Please upload a base resume and enter a job description.")
        return

    status_var.set("Extracting keywords...")
    root.update()
    keywords = extract_keywords(job_desc)

    status_var.set("Generating rewritten resume...")
    root.update()
    rewritten_resume_text = generate_resume(base_resume, job_desc, keywords)
    update_paragraphs(base_resume.paragraphs, rewritten_resume_text)

    status_var.set("Generating cover letter...")
    root.update()
    cover_letter_text = generate_cover_letter(job_desc, base_resume, keywords)

    status_var.set("Review and edit the generated content, then click 'Save Output' to save the files.")
    root.update()

    rewritten_resume_textbox.delete("1.0", tk.END)
    rewritten_resume_textbox.insert(tk.END, rewritten_resume_text)
    cover_letter_textbox.delete("1.0", tk.END)
    cover_letter_textbox.insert(tk.END, cover_letter_text)

def save_output():
    output_folder = filedialog.askdirectory()
    if output_folder:
        rewritten_resume = Document(base_resume_path.get())
        update_paragraphs(rewritten_resume.paragraphs, rewritten_resume_textbox.get("1.0", tk.END))
        rewritten_resume.save(os.path.join(output_folder, "rewritten_resume.docx"))

        cover_letter = Document()
        for para in cover_letter_textbox.get("1.0", tk.END).split("\n"):
            cover_letter.add_paragraph(para)
        cover_letter.save(os.path.join(output_folder, "cover_letter.docx"))

        messagebox.showinfo("Success", "Files saved successfully.")

root = tk.Tk()
root.title("ATS Optimizer")

base_resume_path = tk.StringVar()

job_description_label = tk.Label(root, text="Job Description:")
job_description_label.grid(row=0, column=0, sticky="w")
job_description_text = tk.Text(root, wrap=tk.WORD, width=40, height=10)
job_description_text.grid(row=1, column=0, padx=5, pady=5)

upload_button = tk.Button(root, text="Upload Base Resume", command=upload_base_resume)
upload_button.grid(row=2, column=0, pady=5)

generate_button = tk.Button(root, text="Generate Output", command=generate_output)
generate_button.grid(row=3, column=0, pady=5)

status_var = tk.StringVar()
status_label = tk.Label(root, textvariable=status_var)
status_label.grid(row=4, column=0, pady=5)

rewritten_resume_label = tk.Label(root, text="Rewritten Resume:")
rewritten_resume_label.grid(row=0, column=1, sticky="w")
rewritten_resume_textbox = tk.Text(root, wrap=tk.WORD, width=40, height=10)
rewritten_resume_textbox.grid(row=1, column=1, padx=5, pady=5)

cover_letter_label = tk.Label(root, text="Cover Letter:")
cover_letter_label.grid(row=0, column=2, sticky="w")
cover_letter_textbox = tk.Text(root, wrap=tk.WORD, width=40, height=10)
cover_letter_textbox.grid(row=1, column=2, padx=5, pady=5)

save_button = tk.Button(root, text="Save Output", command=save_output)
save_button.grid(row=2, column=1, columnspan=2, pady=5)

root.mainloop()
